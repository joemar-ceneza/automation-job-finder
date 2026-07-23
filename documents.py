"""
documents.py
Renders a master resume to the formats an application actually needs:
Markdown, DOCX, and PDF.

Layout is deliberately plain — single column, standard headings, no tables or
text boxes. Applicant tracking systems parse that reliably and mangle anything
clever, so the boring layout is the correct one.
"""
import logging
import os
import re
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from resume_model import MasterResume

# fpdf2 subsets the embedded font through fontTools, which logs a line per
# table at INFO — dozens of "glyf pruned" messages that bury the run's own
# output. The work is routine; only its failures are worth hearing about.
logging.getLogger("fontTools").setLevel(logging.WARNING)

# fpdf2's built-in fonts are Latin-1 only, so a resume containing "·" or an
# em dash raises on output. A system TrueType font avoids transliterating the
# document; these are the usual Windows locations, checked in order.
_TTF_CANDIDATES = (
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\calibri.ttf",
    r"C:\Windows\Fonts\segoeui.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
)

# Used only when no TrueType font is available at all.
_ASCII_FALLBACKS = {"·": "-", "—": "-", "–": "-", "•": "-",
                    "'": "'", "'": "'", '"': '"', '"': '"', "…": "..."}


# ======================================================
# INTERNAL HELPERS
# ======================================================
def _ensure_parent(path: str) -> None:
    directory = os.path.dirname(path)
    if directory:
        os.makedirs(directory, exist_ok=True)


def _find_font() -> str | None:
    return next((path for path in _TTF_CANDIDATES if os.path.exists(path)),
                None)


def _to_latin1(text: str) -> str:
    """Last-resort transliteration for PDFs with no usable TrueType font."""
    for source, replacement in _ASCII_FALLBACKS.items():
        text = text.replace(source, replacement)
    normalised = unicodedata.normalize("NFKD", text)
    return normalised.encode("latin-1", "ignore").decode("latin-1")


# ======================================================
# MARKDOWN
# ======================================================
def write_markdown(resume: MasterResume, path: str) -> str:
    """Writes the resume as Markdown — the same shape as the master file."""
    _ensure_parent(path)
    with open(path, "w", encoding="utf-8") as handle:
        handle.write(resume.to_markdown())
    logging.info("Wrote Markdown resume to %s", path)
    return path


# ======================================================
# DOCX
# ======================================================
def _docx_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)


def write_docx(resume: MasterResume, path: str) -> str:
    """
    Writes a single-column DOCX. No tables, no columns, no headers or footers:
    every one of those is a known way to confuse a resume parser.
    """
    _ensure_parent(path)
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.add_run(resume.contact.name)
    name_run.bold = True
    name_run.font.size = Pt(18)

    for line in (resume.contact.headline, resume.contact.detail_line()):
        if not line:
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(line).font.size = Pt(9.5)

    for section in resume.sections:
        _docx_heading(document, section.name)
        if section.prose:
            document.add_paragraph(section.prose)
        if section.items:
            document.add_paragraph(", ".join(section.items))
        for entry in section.entries:
            if entry.title or entry.organisation:
                heading = document.add_paragraph()
                heading.paragraph_format.space_before = Pt(6)
                heading.paragraph_format.space_after = Pt(0)
                title_run = heading.add_run(entry.title)
                title_run.bold = True
                if entry.organisation:
                    heading.add_run(f" — {entry.organisation}")
            if entry.meta:
                meta = document.add_paragraph()
                meta.paragraph_format.space_after = Pt(2)
                meta_run = meta.add_run(entry.meta)
                meta_run.italic = True
                meta_run.font.size = Pt(9.5)
            for bullet in entry.bullets:
                document.add_paragraph(bullet, style="List Bullet")

    document.save(path)
    logging.info("Wrote DOCX resume to %s", path)
    return path


# ======================================================
# PDF
# ======================================================
class _ResumePDF(FPDF):
    """A plain single-column PDF with selectable text."""

    def __init__(self) -> None:
        super().__init__(format="A4", unit="mm")
        self.set_auto_page_break(auto=True, margin=15)
        self.set_margins(18, 15, 18)
        font_path = _find_font()
        self.unicode_ok = font_path is not None
        if self.unicode_ok:
            self.add_font("body", "", font_path)
            self.add_font("body", "B", font_path)
            self.family_name = "body"
        else:
            logging.warning("No TrueType font found — the PDF will fall back "
                            "to Latin-1 and drop unsupported characters.")
            self.family_name = "helvetica"

    def clean(self, text: str) -> str:
        return text if self.unicode_ok else _to_latin1(text)

    def block(self, text: str, size: float, style: str = "",
              gap: float = 4.6, align: str = "L") -> None:
        """
        Writes one wrapped block and returns the cursor to the left margin.

        multi_cell leaves x at the right edge of the cell it just drew, so a
        following call finds no usable width and fpdf2 raises "Not enough
        horizontal space to render a single character". Every block must reset
        the position explicitly.
        """
        self.set_font(self.family_name, style, size)
        self.set_x(self.l_margin)
        self.multi_cell(0, gap, self.clean(text), align=align,
                        new_x=XPos.LMARGIN, new_y=YPos.NEXT)


def write_pdf(resume: MasterResume, path: str) -> str:
    """Writes a single-column PDF with real, selectable text."""
    _ensure_parent(path)
    pdf = _ResumePDF()
    pdf.add_page()

    pdf.block(resume.contact.name, 18, style="B", gap=8, align="C")
    for line in (resume.contact.headline, resume.contact.detail_line()):
        if line:
            pdf.block(line, 9.5, align="C")
    pdf.ln(3)

    for section in resume.sections:
        pdf.block(section.name.upper(), 11, style="B", gap=5.5)
        pdf.set_draw_color(150, 150, 150)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(1.5)

        if section.prose:
            pdf.block(section.prose, 10)
        if section.items:
            pdf.block(", ".join(section.items), 10)

        for entry in section.entries:
            if entry.title or entry.organisation:
                pdf.ln(1)
                pdf.block(entry.heading(), 10.5, style="B")
            if entry.meta:
                pdf.block(entry.meta, 9)
            for bullet in entry.bullets:
                marker = "•" if pdf.unicode_ok else "-"
                pdf.block(f"{marker}  {bullet}", 10)
        pdf.ln(2)

    pdf.output(path)
    logging.info("Wrote PDF resume to %s", path)
    return path


# ======================================================
# COVER LETTERS
# ======================================================
def write_letter_markdown(letter, path: str) -> str:
    """Writes the letter as Markdown."""
    _ensure_parent(path)
    with open(path, "w", encoding="utf-8") as handle:
        handle.write(letter.to_markdown())
    logging.info("Wrote Markdown cover letter to %s", path)
    return path


def write_letter_docx(letter, path: str) -> str:
    """Writes the letter as a plain single-column DOCX."""
    _ensure_parent(path)
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in (letter.letter_date, letter.company):
        if line:
            document.add_paragraph(line)
    document.add_paragraph("")
    document.add_paragraph(letter.salutation())
    for paragraph in letter.paragraphs:
        document.add_paragraph(paragraph)
    document.add_paragraph("")
    document.add_paragraph("Sincerely,")
    signature = document.add_paragraph()
    signature.add_run(letter.sender.name).bold = True
    detail = letter.sender.detail_line()
    if detail:
        document.add_paragraph(detail)

    document.save(path)
    logging.info("Wrote DOCX cover letter to %s", path)
    return path


def write_letter_pdf(letter, path: str) -> str:
    """Writes the letter as a PDF with selectable text."""
    _ensure_parent(path)
    pdf = _ResumePDF()
    pdf.add_page()

    for line in (letter.letter_date, letter.company):
        if line:
            pdf.block(line, 10.5)
    pdf.ln(4)
    pdf.block(letter.salutation(), 11)
    pdf.ln(2)
    for paragraph in letter.paragraphs:
        pdf.block(paragraph, 11, gap=5.4)
        pdf.ln(2)
    pdf.ln(3)
    pdf.block("Sincerely,", 11)
    pdf.block(letter.sender.name, 11, style="B")
    detail = letter.sender.detail_line()
    if detail:
        pdf.block(detail, 9.5)

    pdf.output(path)
    logging.info("Wrote PDF cover letter to %s", path)
    return path


# ======================================================
# PUBLIC ENTRY POINT
# ======================================================
_WRITERS = {"md": write_markdown, "markdown": write_markdown,
            "docx": write_docx, "pdf": write_pdf}

_LETTER_WRITERS = {"md": write_letter_markdown,
                   "markdown": write_letter_markdown,
                   "docx": write_letter_docx, "pdf": write_letter_pdf}


def _dispatch(writers: dict, document, path: str, fmt: str | None) -> str:
    chosen = (fmt or os.path.splitext(path)[1].lstrip(".")).lower()
    writer = writers.get(chosen)
    if writer is None:
        raise ValueError(f"Unsupported format {chosen!r}. "
                         f"Choose from: {', '.join(sorted(set(writers)))}")
    return writer(document, path)


def write(resume: MasterResume, path: str, fmt: str | None = None) -> str:
    """Writes the resume in the format implied by the path, or fmt if given."""
    return _dispatch(_WRITERS, resume, path, fmt)


def write_letter(letter, path: str, fmt: str | None = None) -> str:
    """Writes a cover letter in the format implied by the path, or fmt."""
    return _dispatch(_LETTER_WRITERS, letter, path, fmt)


def slugify(text: str, limit: int = 48) -> str:
    """
    Filesystem-safe stem from a company or job title. Untrusted scraped text
    reaches this, so it is whitelisted rather than merely escaped.
    """
    ascii_text = unicodedata.normalize("NFKD", text or "")
    ascii_text = ascii_text.encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^A-Za-z0-9]+", "-", ascii_text).strip("-").lower()
    return (slug[:limit].rstrip("-") or "resume")
