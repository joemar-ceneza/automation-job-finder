"""
Tests for document rendering.

An exported resume is only useful if a machine can read it back, so these
assert on extracted text rather than on the file existing.
"""
import os
import sys

import pdfplumber
import pytest
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import documents
import resume_model

RESUME_MD = """# Jóse Ceneza
Full Stack Developer
jose@example.com · +63 900 000 0000 · Manila

## Skills

Python, React.js, PostgreSQL

## Experience

### Developer — Acme Corp
Manila · Jan 2023 - Present
- Built a scraping pipeline that cut manual work by 30%.
- Automated reporting with Playwright.
"""


@pytest.fixture
def resume():
    return resume_model.parse_markdown(RESUME_MD)


# ======================================================
# MARKDOWN
# ======================================================
def test_markdown_round_trips_through_the_writer(tmp_path, resume):
    path = tmp_path / "out" / "resume.md"
    documents.write(resume, str(path))
    assert resume_model.load(str(path)) == resume


# ======================================================
# DOCX
# ======================================================
def test_docx_contains_the_content(tmp_path, resume):
    path = tmp_path / "resume.docx"
    documents.write(resume, str(path))
    text = "\n".join(p.text for p in Document(str(path)).paragraphs)
    assert "Jóse Ceneza" in text
    assert "jose@example.com" in text
    assert "Acme Corp" in text
    assert "cut manual work by 30%" in text


def test_docx_has_no_tables(tmp_path, resume):
    """Tables are a known way to confuse resume parsers."""
    path = tmp_path / "resume.docx"
    documents.write(resume, str(path))
    assert Document(str(path)).tables == []


# ======================================================
# PDF
# ======================================================
def test_pdf_text_is_extractable(tmp_path, resume):
    """The entire point: an ATS must be able to read it."""
    path = tmp_path / "resume.pdf"
    documents.write(resume, str(path))
    with pdfplumber.open(str(path)) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    assert "jose@example.com" in text
    assert "Acme Corp" in text
    assert "Playwright" in text


def test_pdf_handles_non_latin1_characters(tmp_path, resume):
    """
    Regression: fpdf2's built-in fonts are Latin-1 only, so an accent or a
    middot raised on output until a TrueType font was loaded.
    """
    path = tmp_path / "resume.pdf"
    documents.write(resume, str(path))       # name has an accent, contact a ·
    assert path.stat().st_size > 1000


def test_pdf_does_not_run_out_of_horizontal_space(tmp_path):
    """
    Regression: multi_cell leaves x at the right edge, so consecutive centred
    blocks raised "Not enough horizontal space to render a single character".
    """
    long_contact = resume_model.parse_markdown(
        "# A Very Long Candidate Name Indeed\n"
        "A Long Professional Headline That Wraps\n"
        "someone.with.a.long.address@example.com · +63 900 000 0000 · "
        "Quezon City, Metro Manila, Philippines · github.com/someone\n"
    )
    path = tmp_path / "resume.pdf"
    documents.write(long_contact, str(path))   # must not raise
    assert path.exists()


# ======================================================
# DISPATCH AND NAMING
# ======================================================
def test_format_is_inferred_from_the_extension(tmp_path, resume):
    for extension in ("md", "docx", "pdf"):
        path = tmp_path / f"resume.{extension}"
        assert documents.write(resume, str(path)) == str(path)
        assert path.exists()


def test_unknown_format_is_rejected(tmp_path, resume):
    with pytest.raises(ValueError, match="Unsupported format"):
        documents.write(resume, str(tmp_path / "resume.rtf"))


@pytest.mark.parametrize("raw, expected", [
    ("Senior Python Developer", "senior-python-developer"),
    ("R&D / Analytics", "r-d-analytics"),
    # Accents transliterate to their base letter rather than being dropped,
    # so an accented name stays readable in the filename.
    ("Café Münchén", "cafe-munchen"),
    ("../../etc/passwd", "etc-passwd"),
    ("", "resume"),
    ("!!!", "resume"),
])
def test_slugify_is_filesystem_safe(raw, expected):
    """Scraped job titles reach this, so it whitelists rather than escapes."""
    slug = documents.slugify(raw)
    assert slug == expected
    assert "/" not in slug and "\\" not in slug and ".." not in slug
